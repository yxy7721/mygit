# -*- coding: utf-8 -*-
"""
Created on Wed Sep  7 11:07:59 2022

@author: yangxy
"""

import pandas as pd
import numpy as np
import docx
import os
import xlwings as xw

app=xw.App(visible=False,add_book=False)
app.display_alerts=False
app.screen_updating=False

docxpath=r"D:\desktop"
docxname=r"交易部周报(20240102-20240105).docx"
destname=r"D:\desktop\交易部周报new.docx"
startdate=pd.to_datetime('2024.1.8',format='%Y.%m.%d')
enddate=pd.to_datetime('2024.1.12',format='%Y.%m.%d')
beixiang=-24.4


excelspath=r"D:\desktop\zhoubao"
xlslist=os.listdir(excelspath)
doc=docx.Document(os.path.join(docxpath,docxname))

'''
print(doc.paragraphs)
print(doc.paragraphs[0].runs[0].text)
doc.paragraphs[0].runs[0].text=r"测试测试"
len(doc.paragraphs)
'''

paraflag=0

#开始摘要第一段市场情况
while True:
    if "权益部分" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1

wb=app.books.open(r"D:\desktop\周报取数t.xlsx")
tmp1=wb.sheets[0].range('A1','F7').options(pd.DataFrame,index=False).value
tmp2=wb.sheets[1].range('B18','D48').options(pd.DataFrame,index=False,expand='table').value
xifen=wb.sheets[2].range('A1').options(pd.DataFrame,index=False,expand='table').value
wb.close()

for i in range(len(doc.paragraphs[paraflag].runs)):
    doc.paragraphs[paraflag].runs[i].text
    if "年" in doc.paragraphs[paraflag].runs[i].text:
        break
run2change=doc.paragraphs[paraflag].runs[i]
run2change.text
text=[]
#text2change[:text2change.find("，")]
text.append(startdate.strftime("%Y年%#m月%#d日至")+enddate.strftime("%#m月%#d日")+"，")
tmp3=tmp1[tmp1["指数名称"]=="上证指数"].iat[0,4]
tmp3= "上涨"+"%.2f" % tmp3 if tmp3>0 else "下跌"+"%.2f" % abs(tmp3)
text.append("上证指数"+tmp3+"%，")
tmp3=tmp1[tmp1["指数名称"]=="创业板"].iat[0,4]
tmp3= "上涨"+"%.2f" % tmp3 if tmp3>0 else "下跌"+"%.2f" % abs(tmp3)
text.append("创业板指数"+tmp3+"%，")
tmp3=tmp1[tmp1["指数名称"]=="沪深300"].iat[0,4]
tmp3= "上涨"+"%.2f" % tmp3 if tmp3>0 else "下跌"+"%.2f" % abs(tmp3)
text.append("沪深300"+tmp3+"%，")
tmp3=tmp1[tmp1["指数名称"]=="中证500"].iat[0,4]
tmp3= "上涨"+"%.2f" % tmp3 if tmp3>0 else "下跌"+"%.2f" % abs(tmp3)
text.append("中证500"+tmp3+"%。")
tmp3=tmp1.iat[4,5]/1
tmp3="%.2f" % (tmp3/10000)
text.append("本周市场日均成交量"+tmp3+"万亿。")
tmp3=beixiang
tmp3="净流入"+"%.2f" % tmp3 if tmp3>0 else "净流出"+"%.2f" % abs(tmp3)
text.append("北向资金"+tmp3+"亿元。")
tmp3=pd.DataFrame(tmp2.columns).T
tmp2.columns=[0,1,2]
tmp2=pd.concat([tmp3,tmp2],axis=0)
tmp3="本周"+"".join(list(tmp2.sort_values(by=2,ascending=False).iloc[:3,1])).replace("(申万)","、",2).replace("(申万)","")+"领涨，"
tmp3=tmp3+"".join(list(tmp2.sort_values(by=2,ascending=True).iloc[:3,1])).replace("(申万)","、",2).replace("(申万)","")+"跌幅最大。"
text.append(tmp3)
tmp3=""
doc.paragraphs[paraflag].runs[i].text="".join(text)
for j in range(i+1,len(doc.paragraphs[paraflag].runs)):
    doc.paragraphs[paraflag].runs[j].text=""
nextuse="".join(text)
nextuse1=text[-1]
del text

#开始搞摘要第二段成交金额方面
while True:
    if "买卖操作" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
doc.paragraphs[paraflag].runs[0].text
#doc.paragraphs[paraflag].runs[1].text

tmp3=list(map(lambda x: 'yes' if x.find('交易所指令报表')!=-1 else x,xlslist))
tmp3=xlslist[tmp3.index('yes')]
wb=app.books.open(os.path.join(excelspath,tmp3))
zhilingtable=wb.sheets[0].range(1,1).options(pd.DataFrame,index=False,expand='table').value
wb.close()
zhilingtable=zhilingtable[zhilingtable['证券类别']=="股票"]
zhilingtable=zhilingtable[zhilingtable['委托方向'].map(lambda x:x in "买入卖出" )]
zhilingtable.columns
pivtable=pd.pivot_table(zhilingtable,index="下达人",aggfunc=np.sum,
                        values="累计成交金额",columns=["委托方向"],margins=True,
                        margins_name='合计')
pivtable=pivtable.drop([pivtable.index[-1]],axis=0)
tmp4=pivtable.sort_values(by="合计",ascending=False)
if "于洋" in set(tmp4.iloc[:4,:].index):
    pass
else:
    tmp4=pd.concat([tmp4.iloc[:3,:],tmp4.loc[["于洋"],:]])
text=[]
tmp3="本周各权益产品的买卖操作方面，"+tmp4.index[0]+"本周操作金额最高，总额达"
tmp3=tmp3+"%.2f" % (tmp4.iat[0,2]/100000000) +"亿元。"
text.append(tmp3)
tmp3="其次是"+tmp4.index[1]+"，本周操作金额"
tmp3=tmp3+"%.2f" % (tmp4.iat[1,2]/100000000)+"亿元。"
text.append(tmp3)
tmp3="然后是"+tmp4.index[2]+"，本周操作金额"
tmp3=tmp3+"%.2f" % (tmp4.iat[2,2]/100000000)+"亿元。"
text.append(tmp3)
tmp3="然后是"+tmp4.index[3]+"，本周操作金额"
tmp3=tmp3+"%.2f" % (tmp4.iat[3,2]/100000000)+"亿元。"
text.append(tmp3)
pivtable=pd.pivot_table(zhilingtable,index=["下达人","基金名称","行业"],aggfunc=np.sum,
                        values="累计成交金额",columns="委托方向",margins=True,
                        margins_name='合计')
tmp5=pd.pivot_table(zhilingtable,index=["下达人","基金名称"],aggfunc=np.sum,
                        values="累计成交金额",columns="委托方向",margins=True,
                        margins_name='合计')
for i in range(4):
    tmp6=tmp5.loc[tmp4.index[i],:].sort_values(by="合计",ascending=False)
    if "消费升级一年持有期" in set(tmp6.index):
        if "消费升级一年持有期" in set(tmp6.index[:2]):
            pass
        else:
            tmp6=pd.concat([tmp6.loc[["消费升级一年持有期"],:],tmp6.iloc[:1,:]])
    tmp3=pivtable.loc[tmp4.index[i],:].loc[tmp6.index[0],:].sort_values(by="合计",ascending=False)
    #pivtable.swaplevel('下达人','基金名称').loc['北京人寿1号',:] 这个很牛逼
    tmp3=tmp3.fillna(0)
    if len(tmp3.index)==1:
        tmp7=""
        if tmp3.iat[0,0]/tmp3.iat[0,2]<=0.3:
            tmp7="减仓"
        elif tmp3.iat[0,0]/tmp3.iat[0,2]>=0.6:
            tmp7="加仓"
        else:
            tmp7="调仓"
        tmp3=tmp6.index[0]+"对"+tmp3.index[0]+"板块进行了"+tmp7[0:2]+"。"
        text[i]=text[i]+tmp3
    else:
        tmp7=""
        if tmp3.iat[0,0]/tmp3.iat[0,2]<=0.3:
            tmp7="减仓"
        elif tmp3.iat[0,0]/tmp3.iat[0,2]>=0.6:
            tmp7="加仓"
        else:
            tmp7="调仓"
        if tmp3.iat[1,0]/tmp3.iat[1,2]<=0.3:
            tmp7=tmp7+"减仓"
        elif tmp3.iat[1,0]/tmp3.iat[1,2]>=0.6:
            tmp7=tmp7+"加仓"
        else:
            tmp7=tmp7+"调仓"
        tmp3=tmp6.index[0]+"对"+tmp3.index[0]+"板块进行了"+tmp7[0:2]+"、"+"对"+tmp3.index[1]+"板块进行了"+tmp7[2:]+"。"
        text[i]=text[i]+tmp3
    
    if len(tmp6.index)==1:
        continue
    tmp3=pivtable.loc[tmp4.index[i],:].loc[tmp6.index[1],:].sort_values(by="合计",ascending=False)
    tmp3=tmp3.fillna(0)
    if len(tmp3.index)==1:
        tmp7=""
        if tmp3.iat[0,0]/tmp3.iat[0,2]<=0.3:
            tmp7="减仓"
        elif tmp3.iat[0,0]/tmp3.iat[0,2]>=0.6:
            tmp7="加仓"
        else:
            tmp7="调仓"
        tmp3=tmp6.index[0]+"对"+tmp3.index[0]+"板块进行了"+tmp7[0:2]+"。"
        text[i]=text[i]+tmp3
    else:
        tmp7=""
        if tmp3.iat[0,0]/tmp3.iat[0,2]<=0.3:
            tmp7="减仓"
        elif tmp3.iat[0,0]/tmp3.iat[0,2]>=0.6:
            tmp7="加仓"
        else:
            tmp7="调仓"
        if tmp3.iat[1,0]/tmp3.iat[1,2]<=0.3:
            tmp7=tmp7+"减仓"
        elif tmp3.iat[1,0]/tmp3.iat[1,2]>=0.6:
            tmp7=tmp7+"加仓"
        else:
            tmp7=tmp7+"调仓"
        tmp3=tmp6.index[1]+"对"+tmp3.index[0]+"板块进行了"+tmp7[0:2]+"、"+"对"+tmp3.index[1]+"板块进行了"+tmp7[2:]+"。"
        text[i]=text[i]+tmp3
        
tmp7=""
del tmp3,tmp4,tmp5,tmp6,tmp7,zhilingtable,pivtable,run2change
doc.paragraphs[paraflag].runs[0].text="".join(text)
for i in range(2,len(doc.paragraphs[paraflag].runs)):
    doc.paragraphs[paraflag].runs[i].text=""
del text

#开始搞换手率那一段
while True:
    if "换手率" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
doc.paragraphs[paraflag].runs[0].text
#doc.paragraphs[paraflag].runs[1].text

tmp3=list(map(lambda x: 'yes' if x.find('换手率查询new')!=-1 else x,xlslist))
tmp3=xlslist[tmp3.index('yes')]
wb=app.books.open(os.path.join(excelspath,tmp3))
huanshou=wb.sheets[0].range(1,1).options(pd.DataFrame,index=False,expand='table').value
wb.close()
for i in range(len(huanshou.index)):
    if huanshou.iat[i,0]=="中证有色金属指数发起式":
        huanshou.iat[i,0]="中证有色"
    elif huanshou.iat[i,0]=="智选核心回报6个月持有期":
        huanshou.iat[i,0]="智选核心回报"
    elif huanshou.iat[i,0]=="稳兴增益六个月持有期":
        huanshou.iat[i,0]="稳兴增益"
    elif huanshou.iat[i,0]=="新能源汽车型发起式":
        huanshou.iat[i,0]="新能源车"
    elif huanshou.iat[i,0]=="数字经济发起式":
        huanshou.iat[i,0]="数字经济"
    elif huanshou.iat[i,0]=="品质消费发起式":
        huanshou.iat[i,0]="品质消费"
    elif huanshou.iat[i,0]=="新添益6个月持有期":
        huanshou.iat[i,0]="新添益"
    elif huanshou.iat[i,0]=="价值混合":
        huanshou.iat[i,0]="价值成长"
    elif huanshou.iat[i,0]=="消费精选混合":
        huanshou.iat[i,0]="消费精选"
    elif huanshou.iat[i,0]=="鑫逸混合":
        huanshou.iat[i,0]="鑫逸"
    elif huanshou.iat[i,0]=="中证钢铁指数发起式" or huanshou.iat[i,0]=="新聚益6个月持有期" or huanshou.iat[i,0]=="财运1号"  or huanshou.iat[i,0]=="臻选1号":
        huanshou.iat[i,0]="del"
    elif huanshou.iat[i,0]=="北京人寿1号" or huanshou.iat[i,0]=="财通500":
        huanshou.iat[i,0]="del"
    elif huanshou.iat[i,0]=="智选FOF16号" or huanshou.iat[i,0]=="智选FOF15号" or huanshou.iat[i,0]=="财兴多策略6号":
        huanshou.iat[i,0]="del"
    elif huanshou.iat[i,0]=="鑫锐":
        huanshou.iat[i,11]="辛晨晨"
    elif huanshou.iat[i,0]=="鑫逸":
        huanshou.iat[i,11]="李晶"
    elif huanshou.iat[i,0]=="科技创新":
        huanshou.iat[i,11]="包斅文"
huanshou=huanshou.drop(huanshou[huanshou['资产单元']=="del"].index,axis=0)
huanshou=huanshou.reset_index(drop=True)

tmp3=list(map(lambda x: 'yes' if x.find('换手率查询old')!=-1 else x,xlslist))
tmp3=xlslist[tmp3.index('yes')]
wb=app.books.open(os.path.join(excelspath,tmp3))
huanshou_old=wb.sheets[0].range(1,1).options(pd.DataFrame,index=False,expand='table').value
wb.close()
for i in range(len(huanshou_old.index)):
    if huanshou_old.iat[i,0]=="中证有色金属指数发起式":
        huanshou_old.iat[i,0]="中证有色"
    elif huanshou_old.iat[i,0]=="智选核心回报6个月持有期":
        huanshou_old.iat[i,0]="智选核心回报"
    elif huanshou_old.iat[i,0]=="稳兴增益六个月持有期":
        huanshou_old.iat[i,0]="稳兴增益"
    elif huanshou_old.iat[i,0]=="新能源汽车型发起式":
        huanshou_old.iat[i,0]="新能源车"
    elif huanshou_old.iat[i,0]=="数字经济发起式":
        huanshou_old.iat[i,0]="数字经济"
    elif huanshou_old.iat[i,0]=="品质消费发起式":
        huanshou_old.iat[i,0]="品质消费"
    elif huanshou_old.iat[i,0]=="新添益6个月持有期":
        huanshou_old.iat[i,0]="新添益"
    elif huanshou_old.iat[i,0]=="价值混合":
        huanshou_old.iat[i,0]="价值成长"
    elif huanshou_old.iat[i,0]=="消费精选混合":
        huanshou_old.iat[i,0]="消费精选"
    elif huanshou_old.iat[i,0]=="鑫逸混合":
        huanshou_old.iat[i,0]="鑫逸"
    elif huanshou_old.iat[i,0]=="中证钢铁指数发起式" or huanshou_old.iat[i,0]=="新聚益6个月持有期" or huanshou_old.iat[i,0]=="财运1号"  or huanshou_old.iat[i,0]=="臻选1号":
        huanshou_old.iat[i,0]="del"
    elif huanshou_old.iat[i,0]=="智选FOF16号" or huanshou_old.iat[i,0]=="智选FOF15号" or huanshou_old.iat[i,0]=="财兴多策略6号":
        huanshou_old.iat[i,0]="del"
    elif huanshou_old.iat[i,0]=="北京人寿1号" or huanshou_old.iat[i,0]=="财通500":
        huanshou_old.iat[i,0]="del"
    elif huanshou_old.iat[i,0]=="鑫锐":
        huanshou_old.iat[i,11]="辛晨晨"
    elif huanshou_old.iat[i,0]=="鑫逸":
        huanshou_old.iat[i,11]="李晶"
    elif huanshou_old.iat[i,0]=="科技创新":
        huanshou_old.iat[i,11]="包斅文"
huanshou_old=huanshou_old.drop(huanshou_old[huanshou_old['资产单元']=="del"].index,axis=0)
huanshou_old=huanshou_old.reset_index(drop=True)

tmp4=0
while True:
    if "资产单元" in doc.tables[tmp4].rows[0].cells[0].text:
        break
    tmp4=tmp4+1
tmp5=huanshou_old.copy()
jiubiao=tmp5.copy()

tableflag=tmp4
del tmp3,tmp4,huanshou_old
tmp6=huanshou[["资产单元","结束日仓位占比","区间换手率"]].copy()
tmp5=tmp5[["资产单元","结束日仓位占比","区间换手率"]].copy()
tmp5=pd.merge(tmp5,tmp6,how="outer",on="资产单元",suffixes=("old","new"))


tmp6=tmp5[["结束日仓位占比old","区间换手率old"]]
tmp5=tmp5.drop(["结束日仓位占比old","区间换手率old"],axis=1)
tmp5=pd.concat([tmp5,tmp6],axis=1)


tmp6=pd.DataFrame(tmp5["结束日仓位占比new"]-tmp5["结束日仓位占比old"])
tmp3=pd.DataFrame((tmp5["区间换手率new"]-tmp5["区间换手率old"]))
tmp5=pd.concat([tmp5,tmp6,tmp3],axis=1)
tmp5.columns.values[5]="仓位gap"
tmp5.columns.values[6]="换手gap"
del tmp3,tmp6
tmp3=tmp5[["资产单元","仓位gap"]]
tmp3=pd.concat([tmp3,pd.DataFrame(tmp3.iloc[:,-1]).applymap(lambda x:abs(x))],axis=1)
tmp3.columns=["资产单元","仓位gap","gapabs"]
tmp4=tmp5[["资产单元","换手gap"]]
tmp4=pd.concat([tmp4,pd.DataFrame(tmp4.iloc[:,-1]).applymap(lambda x:abs(x))],axis=1)
tmp4.columns=["资产单元","换手gap","gapabs"]
tmp3.sort_values(by="gapabs",ascending=False)
tmp6=""
for i in range(5):
    chanp=tmp3.sort_values(by="gapabs",ascending=False).iat[i,0]
    c1=tmp5[tmp5["资产单元"]==chanp]["结束日仓位占比new"].iat[0]
    c2=tmp5[tmp5["资产单元"]==chanp]["结束日仓位占比old"].iat[0]
    c3=tmp5[tmp5["资产单元"]==chanp]["仓位gap"].iat[0]
    c3="增加到" if c3>0 else "减少到"
    tmp6=tmp6+chanp+"的仓位从"+"%.2f" % (c2*100)+"%"+c3+"%.2f" % (c1*100)+"%"
    tmp6=tmp6+"，" if i!=4 else tmp6+"。"
nextuse2=tmp6
for i in range(5):
    chanp=tmp4.sort_values(by="gapabs",ascending=False).iat[i,0]
    c1=tmp5[tmp5["资产单元"]==chanp]["区间换手率new"].iat[0]
    c2=tmp5[tmp5["资产单元"]==chanp]["区间换手率old"].iat[0]
    c3=tmp5[tmp5["资产单元"]==chanp]["换手gap"].iat[0]
    c3="增加到" if c3>0 else "减少到"
    tmp6=tmp6+chanp+"的换手率从"+"%.2f" % (c2*100)+"%"+c3+"%.2f" % (c1*100)+"%"
    tmp6=tmp6+"，" if i!=4 else tmp6+"。"
nextuse3=tmp6
doc.paragraphs[paraflag].runs[0].text="换手率和仓位方面，本周"+tmp6+"其他产品仓位和换手率变化不大。"
if len(doc.paragraphs[paraflag].runs)>1:
    doc.paragraphs[paraflag].runs[1].text=""
del beixiang,docxname,tmp3,tmp4,tmp5,tmp6,c1,c2,c3,chanp,i

#开始搞摘要里面新股那一句话
while True:
    if "只新股" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
doc.paragraphs[paraflag].runs[0].text
#doc.paragraphs[paraflag].runs[1].text
tmp3=list(map(lambda x: 'yes' if x.find('新股参与情况统计')!=-1 else x,xlslist))
tmp3=xlslist[tmp3.index('yes')]
wb=app.books.open(os.path.join(excelspath,tmp3))
xingucan=wb.sheets[0].range('B2').options(pd.DataFrame,index=False,expand='table').value
wb.close()
del tmp3
canyu=0
zhongqian=0
for i in range(len(xingucan.index)):
    if xingucan.iloc[i,3:].sum()==0:
        pass
    else:
        canyu=canyu+1
        if xingucan.iat[i,4]==0 and xingucan.iat[i,6]==0:
            pass
        else:
            zhongqian=zhongqian+1
tmp6="本周共发布"+str(len(xingucan.index))+"只新股，我司参与"+str(canyu)+"只新股，入围"+str(zhongqian)+"只新股"
nextxingu=tmp6
doc.paragraphs[paraflag].runs[0].text=tmp6
if len(doc.paragraphs[paraflag].runs)>1:
    doc.paragraphs[paraflag].runs[1].text=""
del tmp6,canyu,zhongqian

#更新排名表
tmp3=huanshou.iloc[:22,0:12].copy()
len(doc.tables[tableflag].rows)
tableflag=0
while True:
    if "资产单元" in doc.tables[tableflag].rows[0].cells[0].text:
        break
    tableflag=tableflag+1
#此处全部由excel产生，所以全部注释掉:
'''
for i in range(1,len(doc.tables[tableflag].rows)):
    for j in range(len(doc.tables[tableflag].rows[0].cells)):
        #doc.tables[tableflag].rows[i].cells[j].text=str(tmp3.iat[i-1,j])
        tmp3.iat[i-1,j]
        if j==1:
            tmp4="%.4f" % tmp3.iat[i-1,j]
        elif j==3 or j==4:
            tmp4="%.2f" % (tmp3.iat[i-1,j]*100)+"%"
        elif j==9 or j==10:
            tmp4="%d" % tmp3.iat[i-1,j]
        else:
            tmp4=str(tmp3.iat[i-1,j])
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text=tmp4
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>1:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[1].text=""
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>2:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[2].text=""
del tmp3,tmp4
'''
while True:
    if "产品名称" in doc.tables[tableflag].rows[0].cells[0].text:
        break
    tableflag=tableflag+1
tmp3=huanshou.iloc[23:,0:12].copy()
tmp3=tmp3.drop(tmp3.columns[5:11],axis=1).copy()
for i in range(1,len(doc.tables[tableflag].rows)):
    for j in range(len(doc.tables[tableflag].rows[0].cells)):
        #doc.tables[tableflag].rows[i].cells[j].text=str(tmp3.iat[i-1,j])
        tmp3.iat[i-1,j]
        if j==1:
            tmp4="%.4f" % tmp3.iat[i-1,j]
        elif j==3 or j==4:
            tmp4="%.2f" % (tmp3.iat[i-1,j]*100)+"%"
        else:
            tmp4=str(tmp3.iat[i-1,j])
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text=tmp4
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>1:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[1].text=""
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>2:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[2].text=""
del tmp3,tmp4,jiubiao,huanshou

#开始写正文第一节的段落和表格（指数和北向那啥的）
while True:
    if "权益市场走势" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
paraflag=paraflag+1
doc.paragraphs[paraflag].runs[0].text
#doc.paragraphs[paraflag].runs[1].text
for i in range(len(doc.paragraphs[paraflag].runs)):
    if i==0:
        doc.paragraphs[paraflag].runs[i].text=nextuse
    else:
        doc.paragraphs[paraflag].runs[i].text=""
doc.paragraphs[paraflag].text
del nextuse
while True:
    if "代码" in doc.tables[tableflag].rows[0].cells[1].text:
        break
    tableflag=tableflag+1
tmp1=tmp1.drop([tmp1.columns[-1]],axis=1)
for i in range(1,len(doc.tables[tableflag].rows)):
    for j in range(len(doc.tables[tableflag].rows[0].cells)):
        if j==2:
            tmp3="%.2f" % tmp1.iat[i-1,j]
        elif j==3:
            tmp3="%d" % tmp1.iat[i-1,j]
        elif j==4:
            tmp3="%.2f" % tmp1.iat[i-1,j]
        else:
            tmp3=tmp1.iat[i-1,j]
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text=tmp3
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>1:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[1].text=""
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>2:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[2].text=""
del tmp1,tmp3

#开始写申万行业那一段
while True:
    if "申万行业表现" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
paraflag=paraflag+1
doc.paragraphs[paraflag].runs[0].text=nextuse1
del nextuse1
while True:
    if "在细分" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
doc.paragraphs[paraflag].runs[0].text
tmp1=xifen.sort_values(by="20日",ascending=False).iloc[:6,:]
tmp3="在细分概念板块方面，20个交易日以来"+tmp1.iat[0,1]+"、"+tmp1.iat[1,1]+"、"+tmp1.iat[2,1]+"涨幅明显。"
for i in range(len(doc.paragraphs[paraflag].runs)):
    if i==0:
        doc.paragraphs[paraflag].runs[i].text=tmp3
    else:
        doc.paragraphs[paraflag].runs[i].text=""
del tmp1,tmp3
while True:
    if "板块名称" in doc.tables[tableflag].rows[0].cells[0].text:
        break
    tableflag=tableflag+1
tmp2=tmp2.reset_index(drop=True).sort_values(by=2,ascending=False)
for i in range(len(tmp2.index)):
    tmp1=tmp2.iat[i,1]
    tmp3="%.2f" % tmp2.iat[i,2]
    c2=0 if i<=15 else 2
    c1=i+1 if i<=15 else i+1-16
    doc.tables[tableflag].cell(c1,c2).paragraphs[0].runs[0].text=tmp1
    for j in range(1,len(doc.tables[tableflag].cell(c1,c2).paragraphs[0].runs)):
        doc.tables[tableflag].cell(c1,c2).paragraphs[0].runs[j].text=""
    doc.tables[tableflag].cell(c1,c2+1).paragraphs[0].runs[0].text=tmp3
del tmp1,tmp2,tmp3
tableflag=tableflag+1
tmp1=xifen.sort_values(by="20日",ascending=False).iloc[:6,1:]
for i in range(1,len(doc.tables[tableflag].rows)):
    for j in range(len(doc.tables[tableflag].rows[0].cells)):
        if j==1:
            tmp2="%.2f" % tmp1.iat[i-1,j]
        else:
            tmp2=tmp1.iat[i-1,j]
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text=tmp2
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>1:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[1].text=""
del tmp1,tmp2,xifen

#开始搞换手率那一段
nextuse2
nextuse3=nextuse3.replace(nextuse2,"")
while True:
    if "换手率" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
doc.paragraphs[paraflag].runs[0].text
for i in range(len(doc.paragraphs[paraflag].runs)):
    if i==0:
        doc.paragraphs[paraflag].runs[i].text="换手率："+nextuse3+"其他产品换手率变化不大。"
    else:
        doc.paragraphs[paraflag].runs[i].text=""
del nextuse3
while True:
    if "仓位" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
doc.paragraphs[paraflag].runs[0].text
for i in range(len(doc.paragraphs[paraflag].runs)):
    if i==0:
        doc.paragraphs[paraflag].runs[i].text="仓位："+nextuse2+"其他产品仓位变化不大。"
    else:
        doc.paragraphs[paraflag].runs[i].text=""
del nextuse2,c1,c2

#开始搞归因那一段
while True:
    if "超额收益" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
doc.paragraphs[paraflag].runs[0].text
tmp3=list(map(lambda x: 'yes' if x.find('rinson')!=-1 else x,xlslist))
tmp3=xlslist[tmp3.index('yes')]
wb=app.books.open(os.path.join(excelspath,tmp3))
brisontable=wb.sheets[0].range(1,1).options(pd.DataFrame,index=False,expand='table').value
wb.close()
tmp1=pd.DataFrame(None,columns=brisontable.columns)
for i in range(len(brisontable.index)):
    if brisontable.iat[i,0].find("财通资管")!=-1 or brisontable.iat[i,0]=="财享1号":
        tmp1=pd.concat([tmp1,pd.DataFrame(brisontable.iloc[i,:]).T],axis=0)
del brisontable,tmp3
tmp2={"财通资管消费精选混合":"消费精选","财通资管价值成长混合":"价值成长",
      "财通资管科技创新一年定开混合":"科技创新","财通资管消费升级一年持有期混合":"消费升级",
      "财通资管智选核心回报6个月持有期混合":"智选核心","财通资管健康产业混合":"健康产业",
      "财通资管中证有色金属指数发起式":"中证有色","财通资管新能源汽车混合型发起式":"新能源车"}
tmp1.iloc[:,0]=[tmp2.get(x,x) for x in tmp1.iloc[:,0]]
tmp1=pd.concat([tmp1.iloc[:,0],tmp1.iloc[:,9:14]],axis=1)
tmp1['维度名称']=tmp1['维度名称'].astype('category')
list_custom=["消费精选","价值成长","科技创新","消费升级","中证有色","智选核心","新能源车","健康产业","财享1号"]
tmp1['维度名称'].cat.reorder_categories(list_custom,inplace=True)
tmp1=tmp1.sort_values(by="维度名称")
del list_custom,tmp2
tmp2='上表中我司九只代表性产品超额收益从'+"%.2f" % (tmp1[tmp1.columns[3]].min()*100)+"%~"
tmp2=tmp2+"%.2f" % (tmp1[tmp1.columns[3]].max()*100)+"%不等。"
tmp3=tmp1.sort_values(by=tmp1.columns[3],ascending=False).copy()
tmp2=tmp2+tmp3.iat[0,0]+"表现最优，"
if tmp3.iat[0,5]==tmp3.iloc[:,5].max():
    tmp2=tmp2+"因其选股贡献最高。择时贡献最高的是"+tmp1.sort_values(by=tmp1.columns[4],ascending=False).iat[0,0]+"。"
elif tmp3.iat[0,4]==tmp3.iloc[:,4].max():
    tmp2=tmp2+"因其择时贡献最高。选股贡献最高的是"+tmp1.sort_values(by=tmp1.columns[5],ascending=False).iat[0,0]+"。"
else:
    tmp2+"因其择时贡献和选股贡献综合效应最高。单看选股贡献最高的是"+tmp1.sort_values(by=tmp1.columns[5],ascending=False).iat[0,0]+"，"+"择时贡献最高的是"+tmp1.sort_values(by=tmp1.columns[4],ascending=False).iat[0,0]+"。"
for i in range(len(doc.paragraphs[paraflag].runs)):
    if i==0:
        doc.paragraphs[paraflag].runs[i].text=tmp2
    else:
        doc.paragraphs[paraflag].runs[i].text=""
del tmp2,tmp3
while True:
    if len(doc.tables[tableflag].rows[0].cells)<3:
        pass
    elif "超额" in doc.tables[tableflag].rows[0].cells[3].text:
        break
    tableflag=tableflag+1
for i in range(1,len(doc.tables[tableflag].rows)):
    for j in range(len(doc.tables[tableflag].rows[0].cells)):
        if j==0:
            tmp3=tmp1.iat[i-1,j]
        else:
            tmp3="%.2f" % (tmp1.iat[i-1,j]*100)+"%"
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text=tmp3
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>1:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[1].text=""
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>2:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[2].text=""
del tmp3,tmp1

#搞行业重仓和个股重仓那一段
while True:
    if len(doc.tables[tableflag].rows[0].cells)<3:
        pass
    elif "前五大" in doc.tables[tableflag].rows[0].cells[2].text:
        break
    tableflag=tableflag+1
doc.tables[tableflag].cell(1,2).text
tmp3=list(map(lambda x: 'yes' if x.find('重仓股对比')!=-1 else x,xlslist))
tmp3=xlslist[tmp3.index('yes')]
wb=app.books.open(os.path.join(excelspath,tmp3))
zctable=wb.sheets[0].used_range.options(pd.DataFrame,index=False).value
#zctable=wb.sheets[0].range(1,1).options(pd.DataFrame,index=False,expand='table').value
wb.close()
del tmp3
tmp1=zctable.iloc[:,:7]
tmp2=zctable.iloc[:,8:]
del zctable
#tmp3=np.array(tmp1.columns)
#tmp3[tmp3==None]="del"
#tmp3=np.setdiff1d(tmp3,np.array("del"))
list_custom=["价值成长","财通消费精选","科技创新一年定开混合","消费升级一年持有期","财通资管健康产业混合","财通资管均衡臻选混合","财通资管臻享成长混合","财通资管数字经济混合发起式","中证有色金属指数发起式","新能源汽车","财臻1号","财享1号","财通资管品质消费混合发起式","稳兴增益六个月持有期混合"]
tmp1=tmp1.set_index("产品",drop=True).T.copy()
tmp2=tmp2.set_index("产品",drop=True).T.copy()
for k in range(0,len(list_custom)*3,3):
    tmp4=list(tmp1.columns).index(list_custom[int(k/3)])
    tmp4=tmp1.iloc[:,tmp4:tmp4+3]
    tmp4=tmp4.reset_index(drop=True)
    tmp4=tmp4.drop(tmp4.index[0],axis=0)
    for i in range(k+1,k+4):
        for j in range(2,7):
            if (tmp4.iat[j-2,i-k-1]==None) or tmp4.iat[j-2,i-k-1]=='—':
                tmp5="-"
            elif i-k==1:
                tmp5=tmp4.iat[j-2,i-k-1]
            elif i-k==2:
                if isinstance(tmp4.iat[j-2,i-k-1],str):
                    tmp5=tmp4.iat[j-2,i-k-1]
                else:
                    tmp5="%.2f" % (tmp4.iat[j-2,i-k-1]*100)+"%"
            else:
                if isinstance(tmp4.iat[j-2,i-k-1],str):
                    tmp5=tmp4.iat[j-2,i-k]
                else:
                    tmp5="%.2f" % (tmp4.iat[j-2,i-k-1]*100)+"%"
                
            if i-k==3:
                if tmp5[0]!="-":
                    doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.color.rgb=docx.shared.RGBColor(255,0,0)
                else:
                    doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.color.rgb=docx.shared.RGBColor(0,128,0)                    
            
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text=tmp5

            if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>1:
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[1].text=""
            if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>2:
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[2].text=""

tableflag=tableflag+1
for k in range(0,len(list_custom)*3,3):
    tmp4=list(tmp2.columns).index(list_custom[int(k/3)])
    tmp4=tmp2.iloc[:,tmp4:tmp4+3]
    tmp4=tmp4.reset_index(drop=True)
    tmp4=tmp4.drop(tmp4.index[0],axis=0)
    for i in range(k+1,k+4):
        for j in range(2,7):
            if (tmp4.iat[j-2,i-k-1]==None) or tmp4.iat[j-2,i-k-1]=='—':
                tmp5="-"
            elif i-k==1:
                tmp5=tmp4.iat[j-2,i-k-1]
            elif i-k==2:
                if isinstance(tmp4.iat[j-2,i-k-1],str):
                    tmp5=tmp4.iat[j-2,i-k-1]
                else:
                    tmp5="%.2f" % (tmp4.iat[j-2,i-k-1]*100)+"%"
            else:
                if isinstance(tmp4.iat[j-2,i-k-1],str):
                    tmp5=tmp4.iat[j-2,i-k]
                else:
                    tmp5="%.2f" % (tmp4.iat[j-2,i-k-1]*100)+"%"
                
            if i-k==3:
                if tmp5[0]!="-":
                    doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.color.rgb=docx.shared.RGBColor(255,0,0)
                else:
                    doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.color.rgb=docx.shared.RGBColor(0,128,0)                    
            
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text=tmp5

            if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>1:
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[1].text=""
            if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>2:
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[2].text=""
                
del tmp1,tmp2,tmp4,tmp5,list_custom

#开始更新新股概况那一章
while True:
    if "本周共发布" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
for i in range(len(doc.paragraphs[paraflag].runs)):
    if i==0:
        doc.paragraphs[paraflag].runs[i].text=nextxingu
    else:
        doc.paragraphs[paraflag].runs[i].text=""
del nextxingu
while True:
    if len(doc.tables[tableflag].rows[0].cells)<2:
        pass
    elif "参与公募" in doc.tables[tableflag].rows[0].cells[1].text:
        break
    tableflag=tableflag+1
doc.tables[tableflag].cell(0,2).text
len(doc.tables[tableflag].rows)
len(doc.tables[tableflag].rows[0].cells)
tmp1=0
tmp2=0
for i in range(len(xingucan.index)):
    if xingucan.iat[i,4]==0 and xingucan.iat[i,6]==0:
        pass
    else:
        if xingucan.iat[i,0][-2:]=="SH":
            tmp1=tmp1+1
        else:
            tmp2=tmp2+1
'''
for cell in column.cells: # 遍历列中单元格
    cell._element.getparent().remove(cell._element) # 删除第2列的单元格
'''
if (len(doc.tables[tableflag].rows)-1)>max(tmp1,tmp2):
    gap=(len(doc.tables[tableflag].rows)-1)-max(tmp1,tmp2)
    for i in range(gap):
        doc.tables[tableflag].rows[-1]._element.getparent().remove(doc.tables[tableflag].rows[-1]._element) 
else:
    gap=max(tmp1,tmp2)-(len(doc.tables[tableflag].rows)-1)
    for i in range(gap):
        doc.tables[tableflag].add_row()
        doc.tables[tableflag].rows[-1].height=docx.shared.Cm(1.51)
    for i in range(1,len(doc.tables[tableflag].rows)):
        for j in range(len(doc.tables[tableflag].rows[0].cells)):
            doc.tables[tableflag].cell(i,j).vertical_alignment=docx.enum.table.WD_ALIGN_VERTICAL.CENTER
            doc.tables[tableflag].cell(i,j).paragraphs[0].paragraph_format.alignment=docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
            if i>0 and i%2==0:
                shading_elm_1=docx.oxml.parse_xml(r'<w:shd {} w:fill="D9E2F3"/>'.format(docx.oxml.ns.nsdecls('w')))
                doc.tables[tableflag].rows[i].cells[j]._tc.get_or_add_tcPr().append(shading_elm_1)
tmp3=0
tmp4=0
for i in range(len(xingucan.index)):
    if xingucan.iat[i,4]==0 and xingucan.iat[i,6]==0:
        continue
    else:
        if xingucan.iat[i,0][-2:]=="SH":
            tmp3=tmp3+1
            if len(doc.tables[tableflag].cell(tmp3,0).paragraphs[0].runs)==0:
                doc.tables[tableflag].cell(tmp3,0).paragraphs[0].add_run()
            if len(doc.tables[tableflag].cell(tmp3,1).paragraphs[0].runs)==0:
                doc.tables[tableflag].cell(tmp3,1).paragraphs[0].add_run()
            if len(doc.tables[tableflag].cell(tmp3,2).paragraphs[0].runs)==0:
                doc.tables[tableflag].cell(tmp3,2).paragraphs[0].add_run()
            doc.tables[tableflag].cell(tmp3,0).paragraphs[0].runs[0].text=xingucan.iat[i,1]
            doc.tables[tableflag].cell(tmp3,1).paragraphs[0].runs[0].text='公募产品\n（'+"%d"%xingucan.iat[i,4]+"个）"
            doc.tables[tableflag].cell(tmp3,2).paragraphs[0].runs[0].text='智汇系列产品（'+"%d"%xingucan.iat[i,6]+"个）" if xingucan.iat[i,6]>0 else "-"
            for i in range(1,len(doc.tables[tableflag].cell(tmp3,1).paragraphs[0].runs)):
                doc.tables[tableflag].cell(tmp3,1).paragraphs[0].runs[i].text=""
            for i in range(1,len(doc.tables[tableflag].cell(tmp3,2).paragraphs[0].runs)):
                doc.tables[tableflag].cell(tmp3,2).paragraphs[0].runs[i].text=""
            #doc.tables[tableflag].cell(tmp3,1).paragraphs[0].paragraph_format.first_line_indent=docx.shared.Pt(0)
        else:
            tmp4=tmp4+1
            if len(doc.tables[tableflag].cell(tmp4,3).paragraphs[0].runs)==0:
                doc.tables[tableflag].cell(tmp4,3).paragraphs[0].add_run()
            if len(doc.tables[tableflag].cell(tmp4,4).paragraphs[0].runs)==0:
                doc.tables[tableflag].cell(tmp4,4).paragraphs[0].add_run()
            if len(doc.tables[tableflag].cell(tmp4,5).paragraphs[0].runs)==0:
                doc.tables[tableflag].cell(tmp4,5).paragraphs[0].add_run()
            doc.tables[tableflag].cell(tmp4,3).paragraphs[0].runs[0].text=xingucan.iat[i,1]
            doc.tables[tableflag].cell(tmp4,4).paragraphs[0].runs[0].text='公募产品\n（'+"%d"%xingucan.iat[i,4]+"个）"
            doc.tables[tableflag].cell(tmp4,5).paragraphs[0].runs[0].text='智汇系列产品（'+"%d"%xingucan.iat[i,6]+"个）" if xingucan.iat[i,6]>0 else "-"
            for i in range(1,len(doc.tables[tableflag].cell(tmp4,4).paragraphs[0].runs)):
                doc.tables[tableflag].cell(tmp4,4).paragraphs[0].runs[i].text=""
            for i in range(1,len(doc.tables[tableflag].cell(tmp4,5).paragraphs[0].runs)):
                doc.tables[tableflag].cell(tmp4,5).paragraphs[0].runs[i].text=""
if tmp3<tmp4:
    for i in range(tmp3+1,len(doc.tables[tableflag].rows)):
        for j in range(0,3):
            if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>0:
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text=""
            if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>1:
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[1].text=""
            if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>2:
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[2].text=""
elif tmp3>tmp4:
    for i in range(tmp4+1,len(doc.tables[tableflag].rows)):
        for j in range(3,6):
            if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>0:
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text=""
            if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>1:
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[1].text=""
            if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>2:
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[2].text=""
del tmp1,tmp2,tmp3,tmp4
for i in range(1,len(doc.tables[tableflag].rows)):
    for j in range(len(doc.tables[tableflag].rows[0].cells)):
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)==0:
            continue
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.name="楷体"
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0]._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), r'楷体')
        if (j==0 or j==3)and(len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>0):
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.bold=True
        if j==0 or j==3:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.size=docx.shared.Pt(10.5) #即五号（小五为9，六号为12）
        else:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.size=docx.shared.Pt(11)
        if j==1:
            doc.tables[tableflag].cell(i,j).paragraphs[0].paragraph_format.left_indent=docx.shared.Pt(0) #left_indent、right_indent、first_line_indent
            if len(doc.tables[tableflag].cell(i,j).paragraphs)>1:
                doc.tables[tableflag].cell(i,j).paragraphs[1].paragraph_format.first_line_indent=docx.shared.Pt(0)

del xingucan             

#开始搞新股收益那一段
tmp3=list(map(lambda x: 'yes' if x.find('新股交易')!=-1 else x,xlslist))
tmp3=xlslist[tmp3.index('yes')]
wb=app.books.open(os.path.join(excelspath,tmp3))
xgjytable=wb.sheets[0].used_range.options(pd.DataFrame,index=False).value
#zctable=wb.sheets[0].range(1,1).options(pd.DataFrame,index=False,expand='table').value
wb.close()
xgjytable.iloc[:,-2:]=xgjytable.iloc[:,-2:].applymap(lambda x: "公募" if x=="公募基金" else x)
xgjytable.iloc[:,-2:]=xgjytable.iloc[:,-2:].applymap(lambda x: "私募" if x=="定向主动" else x)
xgjytable.iloc[:,-2:]=xgjytable.iloc[:,-2:].applymap(lambda x: "私募" if x=="小集合" else x)
pivtable=pd.pivot_table(xgjytable,index="证券名称",aggfunc=np.sum,
                        values="收益",columns="产品类别",margins=True,
                        margins_name='合计')
pivtable=pivtable.reset_index(drop=False)
pivtable=pd.concat([pivtable.iloc[:-1,:].sort_values(by="合计",ascending=False),pd.DataFrame(pivtable.iloc[-1,:]).T],axis=0).copy()
pivtable=pivtable.reset_index(drop=True).fillna("-")
if "私募" in set(pivtable.columns):
    pass
else:
    pivtable["私募"]=0
    pivtable=pivtable[['证券名称', '公募', "私募" ,'合计' ]].copy()
if "公募" in set(pivtable.columns):
    pass
else:
    pivtable["公募"]=0
    pivtable=pivtable[['证券名称', '公募', "私募" ,'合计' ]].copy()
while True:
    if len(doc.tables[tableflag].rows[0].cells)<2:
        pass
    elif "公募新股" in doc.tables[tableflag].rows[0].cells[1].text:
        break
    tableflag=tableflag+1
len(doc.tables[tableflag].rows)
len(pivtable.index)
if len(doc.tables[tableflag].rows)-1>len(pivtable.index):
    gap=len(doc.tables[tableflag].rows)-1-len(pivtable.index)
    for i in range(gap):
        doc.tables[tableflag].rows[-1]._element.getparent().remove(doc.tables[tableflag].rows[-1]._element) 
else:
    gap=len(pivtable.index)-(len(doc.tables[tableflag].rows)-1)
    for i in range(gap):
        doc.tables[tableflag].add_row()
        doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.04)
    for i in range(1,len(doc.tables[tableflag].rows)):
        for j in range(len(doc.tables[tableflag].rows[0].cells)):
            doc.tables[tableflag].cell(i,j).vertical_alignment=docx.enum.table.WD_ALIGN_VERTICAL.CENTER
            doc.tables[tableflag].cell(i,j).paragraphs[0].paragraph_format.alignment=docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
            if i>0 and i%2==0:
                shading_elm_1=docx.oxml.parse_xml(r'<w:shd {} w:fill="D9E2F3"/>'.format(docx.oxml.ns.nsdecls('w')))
                doc.tables[tableflag].rows[i].cells[j]._tc.get_or_add_tcPr().append(shading_elm_1)
for i in range(1,len(doc.tables[tableflag].rows)):
    for j in range(len(doc.tables[tableflag].rows[0].cells)):
        if j==0:
            tmp3=pivtable.iat[i-1,j]
        else:
            if pivtable.iat[i-1,j]=="-":
                tmp3="-"
            else:
                tmp3=format(pivtable.iat[i-1,j],",.2f")
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)==0:
            doc.tables[tableflag].cell(i,j).paragraphs[0].add_run()
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text=tmp3
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>1:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[1].text=""
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>2:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[2].text=""
for i in range(1,len(doc.tables[tableflag].rows)):
    for j in range(len(doc.tables[tableflag].rows[0].cells)):
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)==0:
            continue
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.name="楷体"
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0]._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), r'楷体')
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.size=docx.shared.Pt(11) #即五号（小五为9，六号为12）

del xgjytable,gap,tmp3

doc.save(destname)

app.kill()






#----------------------------------------------------------
#周五晚上只运行下方东西即可，不需要走上方任何代码
import pandas as pd
import numpy as np
import docx
import os
import xlwings as xw

app=xw.App(visible=False,add_book=False)
app.display_alerts=False
app.screen_updating=False

docxpath=r"E:\desktop"
docxname=r"交易部周报new.docx"
destname=r"E:\desktop\交易部周报newnew.docx"
doc=docx.Document(os.path.join(docxpath,docxname))
wb=app.books.open(r"E:\desktop\【】【】基金排名取数t.xlsx")
fdtable=wb.sheets[0].range('B23','K41').options(pd.DataFrame,index=False).value
wb.close()

paraflag=0
while True:
    if "公募产品市场排名方面" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
doc.paragraphs[paraflag].runs[0].text

tableflag=0
while True:
    if len(doc.tables[tableflag].rows[0].cells)==10:
        if "今年以来" in doc.tables[tableflag].rows[0].cells[6].text:
            break
    tableflag=tableflag+1

fdtable=fdtable.sort_values(by=fdtable.columns[3],ascending=False)
for i in range(1,len(fdtable.index)+1):
    for j in range(len(doc.tables[tableflag].rows[0].cells)):
        if j==0:
            tmp1=str(fdtable.iat[i-1,j])
        elif j==1:
            tmp1="%.2f" % (fdtable.iat[i-1,j])
        elif j==2:
            tmp1=str(fdtable.iat[i-1,j])
        elif j==3:
            tmp1="%.2f" % (fdtable.iat[i-1,j])
        elif j==8:
            tmp1="%.2f" % (fdtable.iat[i-1,j]*100)+"%"
        else:
            tmp1=str(fdtable.iat[i-1,j])
        doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text=tmp1
        if j==5 or j==7:
            if doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].text[-1]=="↓":
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.color.rgb=docx.shared.RGBColor(0,176,80)
            else:
                doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.color.rgb=docx.shared.RGBColor(255,0,0)
        #if j==5 or j==7:
            #doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.size=docx.shared.Pt(8) #（小五为9，六号为12）
        #elif 1<=j and j<=10:
            #doc.tables[tableflag].cell(i,j).paragraphs[0].runs[0].font.size=docx.shared.Pt(9) #（小五为9，六号为12）
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>1:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[1].text=""
        if len(doc.tables[tableflag].cell(i,j).paragraphs[0].runs)>2:
            doc.tables[tableflag].cell(i,j).paragraphs[0].runs[2].text=""

doc.paragraphs[paraflag].text
tmp1=fdtable[fdtable['资产单元']=="价值成长"]['最近一周同类排名'].iat[0]
tmp2=float(tmp1.split("/")[0])
tmp3=float(tmp1.split("/")[1])
tmp1=fdtable[fdtable['资产单元']=="价值成长"]['周排名变动'].iat[0]
tmp1=float(tmp1[:-1]) if tmp1[-1]=="↑" else (-1)*float(tmp1[:-1])
tmp1="上升了"+"%d" % tmp1+"名" if tmp1>=0 else "下降了"+"%d" % abs(tmp1)+"名" 
tmp4="公募产品市场排名方面，姜总的价值成长周排名"+tmp1+"，本周排名"+"%.2f" % (tmp2/tmp3*100)+"%，"
tmp1=fdtable[fdtable['资产单元']=="价值成长"]['年排名百分位'].iat[0]
tmp1="%.2f" % (tmp1*100)+"%"
tmp4=tmp4+"今年以来排名"+tmp1+"。"
tmp1=fdtable[fdtable['资产单元']=="消费精选"]['最近一周同类排名'].iat[0]
tmp2=float(tmp1.split("/")[0])
tmp3=float(tmp1.split("/")[1])
tmp1=fdtable[fdtable['资产单元']=="消费精选"]['周排名变动'].iat[0]
tmp1=float(tmp1[:-1]) if tmp1[-1]=="↑" else (-1)*float(tmp1[:-1])
tmp1="上升了"+"%d" % tmp1+"名" if tmp1>=0 else "下降了"+"%d" % abs(tmp1)+"名" 
tmp4=tmp4+"于洋的消费精选周排名"+tmp1+"，本周排名"+"%.2f" % (tmp2/tmp3*100)+"%，"
tmp1=fdtable[fdtable['资产单元']=="消费精选"]['年排名百分位'].iat[0]
tmp1="%.2f" % (tmp1*100)+"%"
tmp4=tmp4+"今年以来排名"+tmp1+"。"
tmp1=fdtable[fdtable['资产单元']=="消费升级"]['最近一周同类排名'].iat[0]
tmp2=float(tmp1.split("/")[0])
tmp3=float(tmp1.split("/")[1])
tmp1=fdtable[fdtable['资产单元']=="消费升级"]['周排名变动'].iat[0]
tmp1=float(tmp1[:-1]) if tmp1[-1]=="↑" else (-1)*float(tmp1[:-1])
tmp1="上升了"+"%d" % tmp1+"名" if tmp1>=0 else "下降了"+"%d" % abs(tmp1)+"名" 
tmp4=tmp4+"消费升级周排名"+tmp1+"，本周排名"+"%.2f" % (tmp2/tmp3*100)+"%，"
tmp1=fdtable[fdtable['资产单元']=="消费升级"]['年排名百分位'].iat[0]
tmp1="%.2f" % (tmp1*100)+"%"
tmp4=tmp4+"今年以来排名"+tmp1+"。"
tmp1=fdtable[fdtable['资产单元']=="科技创新"]['最近一周同类排名'].iat[0]
tmp2=float(tmp1.split("/")[0])
tmp3=float(tmp1.split("/")[1])
tmp1=fdtable[fdtable['资产单元']=="科技创新"]['周排名变动'].iat[0]
tmp1=float(tmp1[:-1]) if tmp1[-1]=="↑" else (-1)*float(tmp1[:-1])
tmp1="上升了"+"%d" % tmp1+"名" if tmp1>=0 else "下降了"+"%d" % abs(tmp1)+"名" 
tmp4=tmp4+"包斅文的科技创新周排名"+tmp1+"，本周排名"+"%.2f" % (tmp2/tmp3*100)+"%，"
tmp1=fdtable[fdtable['资产单元']=="科技创新"]['年排名百分位'].iat[0]
tmp1="%.2f" % (tmp1*100)+"%"
tmp4=tmp4+"今年以来排名"+tmp1+"。"
tmp1=fdtable[fdtable['资产单元']=="新能源车"]['最近一周同类排名'].iat[0]
tmp2=float(tmp1.split("/")[0])
tmp3=float(tmp1.split("/")[1])
tmp1=fdtable[fdtable['资产单元']=="新能源车"]['周排名变动'].iat[0]
tmp1=float(tmp1[:-1]) if tmp1[-1]=="↑" else (-1)*float(tmp1[:-1])
tmp1="上升了"+"%d" % tmp1+"名" if tmp1>=0 else "下降了"+"%d" % abs(tmp1)+"名" 
tmp4=tmp4+"邵沙锞的新能源车周排名"+tmp1+"，本周排名"+"%.2f" % (tmp2/tmp3*100)+"%，"
tmp1=fdtable[fdtable['资产单元']=="新能源车"]['年排名百分位'].iat[0]
tmp1="%.2f" % (tmp1*100)+"%"
tmp4=tmp4+"今年以来排名"+tmp1+"。"
tmp1=fdtable[fdtable['资产单元']=="健康产业"]['最近一周同类排名'].iat[0]
tmp2=float(tmp1.split("/")[0])
tmp3=float(tmp1.split("/")[1])
tmp1=fdtable[fdtable['资产单元']=="健康产业"]['周排名变动'].iat[0]
tmp1=float(tmp1[:-1]) if tmp1[-1]=="↑" else (-1)*float(tmp1[:-1])
tmp1="上升了"+"%d" % tmp1+"名" if tmp1>=0 else "下降了"+"%d" % abs(tmp1)+"名" 
tmp4=tmp4+"易小金的健康产业周排名"+tmp1+"，本周排名"+"%.2f" % (tmp2/tmp3*100)+"%，"
tmp1=fdtable[fdtable['资产单元']=="健康产业"]['年排名百分位'].iat[0]
tmp1="%.2f" % (tmp1*100)+"%"
tmp4=tmp4+"今年以来排名"+tmp1+"。"
tmp1=fdtable[fdtable['资产单元']=="均衡臻选"]['最近一周同类排名'].iat[0]
tmp2=float(tmp1.split("/")[0])
tmp3=float(tmp1.split("/")[1])
tmp1=fdtable[fdtable['资产单元']=="均衡臻选"]['周排名变动'].iat[0]
tmp1=float(tmp1[:-1]) if tmp1[-1]=="↑" else (-1)*float(tmp1[:-1])
tmp1="上升了"+"%d" % tmp1+"名" if tmp1>=0 else "下降了"+"%d" % abs(tmp1)+"名" 
tmp4=tmp4+"陈李的均衡臻选周排名"+tmp1+"，本周排名"+"%.2f" % (tmp2/tmp3*100)+"%，"
tmp1=fdtable[fdtable['资产单元']=="均衡臻选"]['年排名百分位'].iat[0]
tmp1="%.2f" % (tmp1*100)+"%"
tmp4=tmp4+"今年以来排名"+tmp1+"。"
tmp1=fdtable[fdtable['资产单元']=="品质消费"]['最近一周同类排名'].iat[0]
tmp2=float(tmp1.split("/")[0])
tmp3=float(tmp1.split("/")[1])
tmp1=fdtable[fdtable['资产单元']=="品质消费"]['周排名变动'].iat[0]
tmp1=float(tmp1[:-1]) if tmp1[-1]=="↑" else (-1)*float(tmp1[:-1])
tmp1="上升了"+"%d" % tmp1+"名" if tmp1>=0 else "下降了"+"%d" % abs(tmp1)+"名" 
tmp4=tmp4+"林伟的品质消费周排名"+tmp1+"，本周排名"+"%.2f" % (tmp2/tmp3*100)+"%，"
tmp1=fdtable[fdtable['资产单元']=="品质消费"]['年排名百分位'].iat[0]
tmp1="%.2f" % (tmp1*100)+"%"
tmp4=tmp4+"今年以来排名"+tmp1+"。"

doc.paragraphs[paraflag].runs[0].text=tmp4
for i in range(1,len(doc.paragraphs[paraflag].runs)):
    doc.paragraphs[paraflag].runs[i].text=""
paraflag=paraflag+1
while True:
    if "公募产品市场排名方面" in doc.paragraphs[paraflag].text:
        break
    paraflag=paraflag+1
doc.paragraphs[paraflag].runs[0].text=tmp4
for i in range(1,len(doc.paragraphs[paraflag].runs)):
    doc.paragraphs[paraflag].runs[i].text=""


doc.save(destname)

app.kill()

