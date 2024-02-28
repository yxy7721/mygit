# -*- coding: utf-8 -*-
"""
Created on Wed Jan 31 13:54:53 2024

@author: yangxy
"""

import pandas as pd
import numpy as np
import docx
import os
import xlwings as xw
import copy

def do_it(res):
    docxpath=r"D:\desktop"
    docxname=r"【V4】2023年全年合作机构情况.docx"
    destname=r"D:\desktop\合作机构情况new.docx"
    doc=docx.Document(os.path.join(docxpath,docxname))
    
    tableflag=0
    while True:
        if "时点金额" in doc.tables[tableflag].rows[0].cells[2].text:
            break
        tableflag=tableflag+1
    
    tmp4=copy.deepcopy(res['client_keytime'])
    tmp4=tmp4.sort_values(by="当日成交金额",ascending=False)
    gap=len(doc.tables[tableflag].rows)-2-len(tmp4.index)
    if gap<0:
        doc.tables[tableflag].rows[-1]._element.getparent().remove(doc.tables[tableflag].rows[-1]._element) 
        for i in range(gap):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        doc.tables[tableflag].add_row()
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
    
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[1].paragraphs[0].runs[0].text=tmp4.index[i-1]
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text="%.2f" % (tmp4.iat[i-1,1])

    
    while True:
        if len(doc.tables[tableflag].rows[0].cells)<4:
            tableflag=tableflag+1
            continue
        if "利率债" in doc.tables[tableflag].rows[0].cells[4].text:
            break
        tableflag=tableflag+1
    
    tmp3=copy.deepcopy(res['client_all'])
    tmp3=tmp3.sort_values(by="当日成交金额",ascending=False)
    gap=len(doc.tables[tableflag].rows)-2-len(tmp3.index)
    if gap<0:
        doc.tables[tableflag].rows[-1]._element.getparent().remove(doc.tables[tableflag].rows[-1]._element) 
        for i in range(abs(gap)):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        doc.tables[tableflag].add_row()
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
    
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[1].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[4].paragraphs[0].add_run()        
        doc.tables[tableflag].rows[i].cells[5].paragraphs[0].add_run()       
        doc.tables[tableflag].rows[i].cells[6].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()        
        doc.tables[tableflag].rows[i].cells[1].paragraphs[0].runs[0].text=tmp3.index[i-1]
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text="%.2f" % (tmp3.iat[i-1,1])
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].runs[0].text="%.2f" % (tmp3.iat[i-1,2])
        doc.tables[tableflag].rows[i].cells[4].paragraphs[0].runs[0].text="%.0f" % (tmp3.iat[i-1,3])
        doc.tables[tableflag].rows[i].cells[5].paragraphs[0].runs[0].text="%.0f" % (tmp3.iat[i-1,4])
        doc.tables[tableflag].rows[i].cells[6].paragraphs[0].runs[0].text="%.0f" % (tmp3.iat[i-1,5])
        
        
    tmp5=res["me_all"]
    tmp6=res["me_keytime"]
    


    split=dict()
    for i in tmp6["类型"].unique():
        #print(i)
        tmp7=copy.deepcopy(tmp6[tmp6["类型"]==i])
        split[i]=tmp7

    print(split.keys())
    
    while True:
        if "货币户" in doc.tables[tableflag].rows[1].cells[1].text:
            break
        tableflag=tableflag+1 
        
    tmp7=split["货币户"]
    gap=len(doc.tables[tableflag].rows)-1-len(tmp7.index)
    if gap<0:
        for i in range(abs(gap)):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
            
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].add_run()       
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text=tmp7.index[i-1]
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].runs[0].text="%.2f" % (tmp7.iat[i-1,2])

    
    while True:
        if "利率台产品" in doc.tables[tableflag].rows[1].cells[1].text:
            break
        tableflag=tableflag+1 
    tmp7=split["利率台产品"]
    gap=len(doc.tables[tableflag].rows)-1-len(tmp7.index)
    if gap<0:
        for i in range(abs(gap)):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
            
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].add_run()       
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text=tmp7.index[i-1]
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].runs[0].text="%.2f" % (tmp7.iat[i-1,2])

    while True:
        if "银行定制" in doc.tables[tableflag].rows[1].cells[1].text:
            break
        tableflag=tableflag+1 
        
    tmp7=split["银行定制"]#这里可能报错 那么下面忽略
    gap=len(doc.tables[tableflag].rows)-1-len(tmp7.index)
    if gap<0:
        for i in range(abs(gap)):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
            
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].add_run()       
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text=tmp7.index[i-1]
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].runs[0].text="%.2f" % (tmp7.iat[i-1,2])

    while True:
        if "鸿系列" in doc.tables[tableflag].rows[1].cells[1].text:
            break
        tableflag=tableflag+1 
        
    tmp7=split["鸿系列"]#这里可能报错 那么下面忽略
    gap=len(doc.tables[tableflag].rows)-1-len(tmp7.index)
    if gap<0:
        for i in range(abs(gap)):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
            
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].add_run()       
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text=tmp7.index[i-1]
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].runs[0].text="%.2f" % (tmp7.iat[i-1,2])
 
    while True:
        if "固收+" in doc.tables[tableflag].rows[1].cells[1].text:
            break
        tableflag=tableflag+1 
        
    tmp7=split["固收+"]#这里可能报错 那么下面忽略
    gap=len(doc.tables[tableflag].rows)-1-len(tmp7.index)
    if gap<0:
        for i in range(abs(gap)):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
            
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].add_run()       
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text=tmp7.index[i-1]
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].runs[0].text="%.2f" % (tmp7.iat[i-1,2])


    split=dict()
    for i in tmp5["类型"].unique():
        #print(i)
        tmp7=copy.deepcopy(tmp5[tmp5["类型"]==i])
        split[i]=tmp7

    print(split.keys())
    
    while True:
        if "货币户" in doc.tables[tableflag].rows[1].cells[1].text:
            break
        tableflag=tableflag+1 
        
    tmp7=split["货币户"]
    gap=len(doc.tables[tableflag].rows)-1-len(tmp7.index)
    if gap<0:
        for i in range(abs(gap)):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
            
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].add_run()       
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text=tmp7.index[i-1]
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].runs[0].text="%.2f" % (tmp7.iat[i-1,2])

    
    while True:
        if "利率台产品" in doc.tables[tableflag].rows[1].cells[1].text:
            break
        tableflag=tableflag+1 
    tmp7=split["利率台产品"]
    gap=len(doc.tables[tableflag].rows)-1-len(tmp7.index)
    if gap<0:
        for i in range(abs(gap)):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
            
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].add_run()       
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text=tmp7.index[i-1]
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].runs[0].text="%.2f" % (tmp7.iat[i-1,2])

    while True:
        if "银行定制" in doc.tables[tableflag].rows[1].cells[1].text:
            break
        tableflag=tableflag+1 
        
    tmp7=split["银行定制"]#这里可能报错 那么下面忽略
    gap=len(doc.tables[tableflag].rows)-1-len(tmp7.index)
    if gap<0:
        for i in range(abs(gap)):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
            
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].add_run()       
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text=tmp7.index[i-1]
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].runs[0].text="%.2f" % (tmp7.iat[i-1,2])

    while True:
        if "鸿系列" in doc.tables[tableflag].rows[1].cells[1].text:
            break
        tableflag=tableflag+1 
        
    tmp7=split["鸿系列"]#这里可能报错 那么下面忽略
    gap=len(doc.tables[tableflag].rows)-1-len(tmp7.index)
    if gap<0:
        for i in range(abs(gap)):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
            
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].add_run()       
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text=tmp7.index[i-1]
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].runs[0].text="%.2f" % (tmp7.iat[i-1,2])
 
    while True:
        if "固收+" in doc.tables[tableflag].rows[1].cells[1].text:
            break
        tableflag=tableflag+1 
        
    tmp7=split["固收+"]#这里可能报错 那么下面忽略
    gap=len(doc.tables[tableflag].rows)-1-len(tmp7.index)
    if gap<0:
        for i in range(abs(gap)):
            doc.tables[tableflag].add_row()
            doc.tables[tableflag].rows[-1].height=docx.shared.Cm(0.05)
        
    else:
        for i in range(gap):
            doc.tables[tableflag].rows[-2]._element.getparent().remove(doc.tables[tableflag].rows[-2]._element) 
            
    for i in range(1,len(doc.tables[tableflag].rows)-1):
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].add_run()
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].add_run()       
        doc.tables[tableflag].rows[i].cells[2].paragraphs[0].runs[0].text=tmp7.index[i-1]
        doc.tables[tableflag].rows[i].cells[3].paragraphs[0].runs[0].text="%.2f" % (tmp7.iat[i-1,2])






    doc.save(destname)
    return





























